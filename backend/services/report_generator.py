from docx import Document
from docx.shared import Inches
import os
from typing import Dict, Any
from datetime import datetime

class ReportGenerator:
    def __init__(self):
        self.reports_dir = "reports"
        os.makedirs(self.reports_dir, exist_ok=True)

    async def generate_report(self, analysis_result: Dict[str, Any]) -> str:
        try:
            # Create a new Document
            doc = Document()

            # Add header
            doc.add_heading('Investment Analysis Report', 0)
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

            # Add analysis section
            doc.add_heading('Analysis', level=1)
            doc.add_paragraph(analysis_result['analysis'])

            # Add recommendations section
            doc.add_heading('Recommendations', level=1)
            for recommendation in analysis_result['recommendations']:
                doc.add_paragraph(f'• {recommendation}')

            # Add confidence score
            doc.add_heading('Confidence Score', level=1)
            doc.add_paragraph(f'Analysis Confidence: {analysis_result["confidence"]*100}%')

            # Save the document
            filename = f'investment_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            filepath = os.path.join(self.reports_dir, filename)
            doc.save(filepath)

            # In a real application, you would upload this file to a cloud storage
            # and return a public URL. For this example, we'll return a local path
            return f"/reports/{filename}"

        except Exception as e:
            raise Exception(f"Report generation failed: {str(e)}")

    def _format_currency(self, amount: float) -> str:
        return "${:,.2f}".format(amount)