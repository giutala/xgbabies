from langchain import PromptTemplate, LLMChain
from langchain.agents import Tool, initialize_agent, AgentType
from langchain.llms import OpenAI
import pandas as pd
import numpy as np
import xgboost as xgb
from sklearn.model_selection import train_test_split
from docx import Document
import traceback


# Initialize LLM
llm = OpenAI(temperature=0.7, model="text-davinci-003")


# Utility function for error handling
def safe_run(agent_func, *args, **kwargs):
    try:
        return agent_func(*args, **kwargs)
    except Exception as e:
        print(f"Error in {agent_func.__name__}: {e}")
        print(traceback.format_exc())
        return f"An error occurred while processing: {e}"


# Agent 1: Market Analysis and Technical Review
def market_analysis(idea):
    prompt = PromptTemplate(
        input_variables=["idea"],
        template="Analyze the market and provide technical insights for the following business idea: {idea}"
    )
    chain = LLMChain(llm=llm, prompt=prompt)
    return safe_run(chain.run, idea)


# Agent 2: Competition Analysis
def competition_analysis(financial_data, competitor_data):
    prompt = PromptTemplate(
        input_variables=["financial_data", "competitor_data"],
        template=(
            "Analyze the competition based on the provided financial data of the company ({financial_data}) "
            "and its competitor ({competitor_data}). Identify correlations and bottlenecks."
        )
    )
    chain = LLMChain(llm=llm, prompt=prompt)
    return safe_run(chain.run, financial_data=financial_data, competitor_data=competitor_data)


# Agent 3: Balance Sheet Analysis
def balance_sheet_analysis(balance_sheet, investment_amount):
    liquidity = balance_sheet.get('cash', 0) - investment_amount
    liquidity_status = "No liquidity issues." if liquidity >= 0 else "Insufficient liquidity."

    prompt = PromptTemplate(
        input_variables=["balance_sheet", "investment_amount"],
        template=(
            "Analyze the balance sheet for liquidity. If there is insufficient cash "
            "({investment_amount} needed), evaluate the company's prospects for securing bank debt. "
            "Balance Sheet: {balance_sheet}"
        )
    )
    chain = LLMChain(llm=llm, prompt=prompt)
    return liquidity_status + "\n" + safe_run(chain.run, balance_sheet=balance_sheet, investment_amount=investment_amount)


# Agent 4: Investment Analysis
def investment_analysis(cash_flows, discount_rate):
    try:
        npv = sum(cf / ((1 + discount_rate) ** i) for i, cf in enumerate(cash_flows))
        roi = (sum(cash_flows) - cash_flows[0]) / cash_flows[0]
        payback_period = next(
            (i for i, _ in enumerate(np.cumsum(cash_flows)) if np.cumsum(cash_flows)[i] >= 0), None
        )
        return {"NPV": npv, "ROI": roi, "Payback Period": payback_period}
    except Exception as e:
        print(f"Error in investment_analysis: {e}")
        print(traceback.format_exc())
        return {"error": str(e)}


# Agent 5: Market Prediction with XGBoost
def market_prediction(data):
    try:
        X = data.drop("target", axis=1)
        y = data["target"]
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        model = xgb.XGBRegressor()
        model.fit(X_train, y_train)
        predictions = model.predict(X_test)
        cagr = (predictions[-1] / predictions[0]) ** (1 / len(predictions)) - 1
        return {"Predictions": predictions.tolist(), "CAGR": cagr}
    except Exception as e:
        print(f"Error in market_prediction: {e}")
        print(traceback.format_exc())
        return {"error": str(e)}


# Agent 6: Predict Cash Flows
def predict_cashflows(product_description, market_area, initial_investment):
    try:
        # Step 1: Search Competitive Price
        prompt_price = PromptTemplate(
            input_variables=["product_description"],
            template=(
                "Based on the description: {product_description}, "
                "determine a competitive market price for the product."
            )
        )
        chain_price = LLMChain(llm=llm, prompt=prompt_price)
        competitive_price = float(safe_run(chain_price.run, product_description))
        
        # Step 2: Estimate Market Sizing
        prompt_market_size = PromptTemplate(
            input_variables=["market_area", "product_description"],
            template=(
                "Estimate the potential market size and yearly sales (in units) "
                "for a product described as: {product_description} "
                "in the area: {market_area} for the next 5 years."
            )
        )
        chain_market_size = LLMChain(llm=llm, prompt=prompt_market_size)
        market_sizing = safe_run(chain_market_size.run, market_area=market_area, product_description=product_description)
        estimated_sales = eval(market_sizing)

        # Step 3: Quantify Costs of Implementation
        prompt_costs = PromptTemplate(
            input_variables=["product_description", "initial_investment", "market_area"],
            template=(
                "Given the product: {product_description}, "
                "an initial investment of {initial_investment}, and the market area: {market_area}, "
                "estimate the yearly operational and implementation costs over the next 5 years."
            )
        )
        chain_costs = LLMChain(llm=llm, prompt=prompt_costs)
        yearly_costs = eval(safe_run(chain_costs.run, product_description=product_description, 
                                     initial_investment=initial_investment, market_area=market_area))

        # Step 4: Calculate Cash Flows
        cashflows = [
            (competitive_price * units_sold) - cost
            for units_sold, cost in zip(estimated_sales, yearly_costs)
        ]
        
        return {
            "Competitive Price": competitive_price,
            "Estimated Sales": estimated_sales,
            "Yearly Costs": yearly_costs,
            "Cashflows": cashflows
        }
    
    except Exception as e:
        print(f"Error in predict_cashflows: {e}")
        print(traceback.format_exc())
        return {"error": str(e)}


# Orchestrator: Generate Full Report
def generate_report_with_cashflows(idea, balance_sheet, competitor_data, market_data, cash_flows, discount_rate, investment_amount, product_description, market_area):
    document = Document()
    document.add_heading("Investment Analysis Report", level=1)

    # Step 1: Market Analysis
    market_analysis_result = market_analysis(idea)
    document.add_heading("Market Analysis", level=2)
    document.add_paragraph(market_analysis_result)

    # Step 2: Competition Analysis
    competition_analysis_result = competition_analysis(balance_sheet, competitor_data)
    document.add_heading("Competition Analysis", level=2)
    document.add_paragraph(competition_analysis_result)

    # Step 3: Balance Sheet Analysis
    balance_sheet_result = balance_sheet_analysis(balance_sheet, investment_amount)
    document.add_heading("Balance Sheet Analysis", level=2)
    document.add_paragraph(balance_sheet_result)

    # Step 4: Investment Analysis
    investment_result = investment_analysis(cash_flows, discount_rate)
    document.add_heading("Investment Analysis", level=2)
    document.add_paragraph(str(investment_result))

    # Step 5: Market Prediction
    market_prediction_result = market_prediction(market_data)
    document.add_heading("Market Prediction", level=2)
    document.add_paragraph(str(market_prediction_result))

    # Step 6: Cash Flow Prediction
    cashflow_prediction_result = predict_cashflows(product_description, market_area, investment_amount)
    document.add_heading("Cash Flow Prediction", level=2)
    document.add_paragraph(str(cashflow_prediction_result))

    # Save Report
    output_file = "Investment_Analysis_Report_With_Cashflows.docx"
    document.save(output_file)
    print(f"Report generated: {output_file}")
    return output_file


# Example Input
idea = "A renewable energy startup focusing on solar panel installations."
balance_sheet = {"cash": 50000, "assets": 200000, "liabilities": 150000}
competitor_data = "Competitor financial and market data"
market_data = pd.DataFrame({"feature1": [1, 2, 3], "feature2": [4, 5, 6], "target": [100, 200, 300]})
cash_flows = [-100000, 20000, 30000, 40000, 50000]
discount_rate = 0.1
investment_amount = 100000
product_description = "High-efficiency solar panels for residential use."
market_area = "Urban areas in California with high solar adoption potential."

# Generate Report
generate_report_with_cashflows(idea, balance_sheet, competitor_data, market_data, cash_flows, discount_rate, investment_amount, product_description, market_area)
